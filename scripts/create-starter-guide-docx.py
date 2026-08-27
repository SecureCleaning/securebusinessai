from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "AI_Starter_Guide_Draft.docx"
LOGO = ROOT / "assets" / "secure-business-ai-stacked.png"

NAVY = RGBColor(16, 32, 51)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 100, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "D9E0EA"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 25, NAVY, 0, 10),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.text = "Secure Business AI | AI Starter Guide Draft"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED
    add_page_number(section.footer.add_paragraph())


def add_para(doc, text, style=None, bold=False, italic=False, color=None, size=None, align=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, "CCD6E3")
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    p2.add_run(body)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = NAVY
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.2
            p.add_run(value)
            if widths:
                cell.width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        p.add_run("[ ] ").bold = True
        p.add_run(item)


def add_cover(doc):
    add_para(doc, "Secure Business AI", bold=True, color=BLUE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.55))
    add_para(doc, "AI for Your Business", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(
        doc,
        "What the fuss is about and why you cannot afford to wait",
        bold=True,
        color=NAVY,
        size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=18,
    )
    add_para(
        doc,
        "A plain-English starter guide for small business owners",
        italic=True,
        color=MUTED,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=24,
    )
    add_callout(
        doc,
        "Draft note",
        "This is a working Word draft for review and editing. It is written to become a 30-50 page designed PDF after content approval, layout design, and final examples are added.",
    )
    add_para(doc, "Prepared by Secure Business AI", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_para(doc, "Draft version 0.1", color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Contents", level=1)
    chapters = [
        "Before You Start",
        "1. What AI Actually Is",
        "2. What AI Is Already Doing",
        "3. Why Waiting Can Cost You",
        "4. A Simple Framework for Knowing Where to Start",
        "5. Tools You Can Use Today",
        "6. The Stages of AI Adoption",
        "7. Your AI Opportunity Map",
        "8. Self-Assessment Checklist",
        "Next Steps",
    ]
    for chapter in chapters:
        add_para(doc, chapter)
    doc.add_page_break()


def chapter(doc, title, intro, sections):
    doc.add_heading(title, level=1)
    add_para(doc, intro)
    for heading, paragraphs, bullets in sections:
        doc.add_heading(heading, level=2)
        for paragraph in paragraphs:
            add_para(doc, paragraph)
        if bullets:
            add_bullets(doc, bullets)
    doc.add_page_break()


def build_doc():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)

    chapter(
        doc,
        "Before You Start",
        "AI is no longer only for large companies or technology teams. It is already being used by small businesses to answer questions, write better first drafts, respond faster, reduce admin, and make websites more useful.",
        [
            (
                "Who this guide is for",
                [
                    "This guide is written for practical business owners. You do not need to understand models, prompts, APIs, or software architecture to get value from AI.",
                    "The aim is to help you see where AI can help your business now, where it should wait, and what to avoid so you do not waste money on disconnected tools.",
                ],
                [
                    "You run or manage a small business.",
                    "You want clearer answers before buying AI services.",
                    "You are curious, but you do not want hype.",
                    "You want AI to support revenue, service, admin, or customer response.",
                ],
            ),
            (
                "The simple promise",
                [
                    "AI should make the business easier to run or easier to buy from. If it does not do either of those things, it is probably a distraction.",
                    "A useful AI project should improve one of four things: enquiries, response speed, customer experience, or internal workload.",
                ],
                None,
            ),
        ],
    )

    chapter(
        doc,
        "1. What AI Actually Is",
        "In plain English, AI is software that can understand information, generate useful responses, and help complete tasks that previously needed a person to think through the first draft.",
        [
            (
                "AI is not magic",
                [
                    "AI does not truly understand your business the way you do. It works from patterns, examples, instructions, and the information it is given.",
                    "That is why training matters. The better the source material, the better the output. A chatbot trained on weak website copy will usually give weak answers. A voice receptionist with no service rules will ask poor questions.",
                ],
                None,
            ),
            (
                "Common types of AI you will hear about",
                [
                    "Different AI tools solve different problems. The names can sound complicated, but the business use cases are usually simple.",
                ],
                [
                    "Chat assistants answer typed questions on a website or inside a business system.",
                    "Voice agents answer phone calls, ask questions, and summarise the call.",
                    "Writing assistants help draft emails, proposals, ads, policies, and website copy.",
                    "Workflow automation moves information between tools and reduces manual follow-up.",
                    "Search and answer tools help customers or staff find the right information faster.",
                ],
            ),
            (
                "The most important idea",
                [
                    "AI is most useful when it is connected to a clear job. Do not start with the tool. Start with the business problem.",
                    "A good first AI project might be as simple as capturing missed calls, improving website enquiries, turning FAQs into better answers, or helping staff prepare better responses.",
                ],
                None,
            ),
        ],
    )

    chapter(
        doc,
        "2. What AI Is Already Doing",
        "Many businesses are already using AI without making a big announcement about it. The competitive advantage often comes from small improvements repeated every day.",
        [
            (
                "Service businesses",
                [
                    "A local service business can use AI to answer common questions, collect job details, explain service areas, capture after-hours enquiries, and prepare cleaner handovers for staff.",
                    "This matters because many customers do not wait. If one business answers clearly and another business says 'call us', the clearer business often wins the enquiry.",
                ],
                [
                    "Answering website questions before the customer calls.",
                    "Collecting quote details after hours.",
                    "Summarising calls so staff do not have to replay or remember details.",
                    "Following up with leads who asked for help but did not book.",
                ],
            ),
            (
                "Professional services",
                [
                    "Professional services firms can use AI to prepare first drafts, intake clients, summarise meetings, maintain knowledge bases, and help prospects understand what happens next.",
                    "The key is not replacing professional judgement. The key is reducing repetitive explanation and admin so the team can spend more time on high-value work.",
                ],
                [
                    "Client intake questionnaires.",
                    "Plain-English service explanations.",
                    "Proposal and email first drafts.",
                    "Meeting summaries and action lists.",
                ],
            ),
            (
                "Retail and local operators",
                [
                    "Retail, hospitality, and local operators can use AI to answer opening hours, product questions, bookings, stock questions, event details, and local recommendations.",
                    "Small improvements in response speed can turn casual interest into a sale or booking.",
                ],
                [
                    "Website chat for common questions.",
                    "Promotion and social content drafts.",
                    "Customer service response templates.",
                    "Staff knowledge guides.",
                ],
            ),
        ],
    )

    doc.add_heading("Industry Example Map", level=2)
    add_table(
        doc,
        ["Industry", "Useful AI starting point", "Why it helps"],
        [
            ("Trades", "Voice receptionist and quote intake", "Captures jobs when the team is on the tools."),
            ("Cleaning", "Website chat plus service area answers", "Helps customers self-qualify and enquire faster."),
            ("Dental/health", "Appointment questions and call summaries", "Reduces repetitive reception workload."),
            ("Accounting", "Client intake and document reminders", "Cuts admin before appointments."),
            ("Retail", "Product questions and campaign drafts", "Helps customers get answers outside staff hours."),
            ("Real estate", "Buyer/seller FAQs and inspection follow-up", "Improves lead response consistency."),
        ],
        widths=[1.35, 2.35, 2.8],
    )
    doc.add_page_break()

    chapter(
        doc,
        "3. Why Waiting Can Cost You",
        "The risk is not that AI will suddenly replace every business. The risk is that competitors use it to respond faster, explain better, and follow up more consistently.",
        [
            (
                "Customers compare speed and clarity",
                [
                    "Most customers are not judging your technology. They are judging how easy it is to understand your offer and get help.",
                    "If AI helps a competitor answer questions, capture details, and follow up quickly, the customer may never reach your business at all.",
                ],
                None,
            ),
            (
                "The cost of missed moments",
                [
                    "Small businesses often lose opportunities in small moments: an unanswered phone call, a vague website, a form that asks too much, or a delayed reply.",
                    "AI can help cover those gaps, especially when staff are busy, unavailable, or repeating the same answers every day.",
                ],
                [
                    "Missed calls after hours.",
                    "Customers who leave because the website does not answer their question.",
                    "Slow follow-up after a quote request.",
                    "Staff time spent rewriting the same explanation.",
                    "No system for turning enquiries into organised next steps.",
                ],
            ),
        ],
    )

    chapter(
        doc,
        "4. A Simple Framework for Knowing Where to Start",
        "The best place to start is usually where the business is already leaking time, leads, or clarity.",
        [
            (
                "The four-question filter",
                [
                    "Before choosing a tool, answer four questions. These questions keep the project tied to business value instead of novelty.",
                ],
                [
                    "Where are we losing enquiries or sales opportunities?",
                    "Where are staff repeating the same answers or admin?",
                    "Where do customers get confused before buying?",
                    "Where would faster response create measurable value?",
                ],
            ),
            (
                "Pick one narrow first project",
                [
                    "Do not try to automate the whole business in one step. The first project should be narrow enough to explain in one sentence.",
                    "For example: 'answer missed calls after hours', 'help website visitors understand our services', or 'summarise quote calls for the office team'.",
                ],
                None,
            ),
        ],
    )

    add_callout(
        doc,
        "Good first project test",
        "If the project cannot be explained simply, measured roughly, and improved over time, make it smaller before you buy software.",
    )
    add_table(
        doc,
        ["If the problem is...", "Start with...", "Avoid starting with..."],
        [
            ("Missed phone calls", "AI voice receptionist", "A full website rebuild before fixing response."),
            ("Poor website enquiries", "AI website review", "Adding a chatbot to unclear pages."),
            ("Repeated customer questions", "Website AI assistant", "Complex back-office automation."),
            ("Admin overload", "Workflow mapping and starter pack", "Buying random AI subscriptions."),
            ("Weak online clarity", "Website messaging review", "Paid ads before fixing the offer."),
        ],
        widths=[2.05, 2.1, 2.35],
    )
    doc.add_page_break()

    chapter(
        doc,
        "5. Tools You Can Use Today",
        "You do not need to buy a large AI system to begin. Many businesses can get value from everyday tools once they use them with clear guardrails.",
        [
            (
                "Everyday AI tools",
                [
                    "General AI assistants can help write first drafts, summarise notes, plan content, explain technical topics, and prepare checklists.",
                    "The important habit is to review the output. AI is a draft assistant, not a final decision maker.",
                ],
                [
                    "Drafting emails and customer responses.",
                    "Turning meeting notes into action items.",
                    "Rewriting technical language into customer language.",
                    "Brainstorming FAQ answers.",
                    "Preparing social posts or website outlines.",
                ],
            ),
            (
                "Business-specific AI tools",
                [
                    "Some AI tools are designed for a specific business job. These are often more useful than general tools because they connect to a workflow.",
                    "Examples include website chatbots, voice receptionists, booking assistants, CRM automation, report generation, and customer support tools.",
                ],
                None,
            ),
            (
                "Use guardrails",
                [
                    "AI should not be allowed to invent prices, promise things your business cannot deliver, or answer sensitive questions without rules.",
                    "Good guardrails include approved answers, escalation rules, privacy boundaries, service areas, booking conditions, and human handover points.",
                ],
                None,
            ),
        ],
    )

    chapter(
        doc,
        "6. The Stages of AI Adoption",
        "AI adoption usually works best in stages. Each stage builds confidence and creates better information for the next stage.",
        [
            (
                "Stage 1: Learn and map",
                [
                    "The business learns what AI can do and maps the biggest leaks in its current customer journey or internal workflow.",
                    "This stage should produce a short list of practical opportunities, not a giant software shopping list.",
                ],
                None,
            ),
            (
                "Stage 2: Fix the front door",
                [
                    "The front door is where customers first decide whether to trust you. It includes the website, phone experience, enquiry forms, service explanations, and follow-up.",
                    "For many small businesses, this is the highest-value place to begin.",
                ],
                None,
            ),
            (
                "Stage 3: Add assistants",
                [
                    "Once the offer and process are clearer, AI assistants can answer website questions, capture calls, collect quote details, and hand information to the team.",
                ],
                None,
            ),
            (
                "Stage 4: Connect workflows",
                [
                    "The next stage is connecting information between systems so customer details, summaries, reminders, and follow-ups move with less manual effort.",
                ],
                None,
            ),
            (
                "Stage 5: Improve over time",
                [
                    "AI is not a set-and-forget project. The best systems improve from better answers, clearer service rules, and real customer questions.",
                ],
                None,
            ),
        ],
    )

    add_table(
        doc,
        ["Stage", "Business question", "Typical output"],
        [
            ("Learn", "What should we understand first?", "Plain-English AI guide and opportunity list."),
            ("Review", "Where are we leaking leads or time?", "Website and AI readiness review."),
            ("Assist", "What should AI answer or capture?", "Chatbot, voice receptionist, or intake assistant."),
            ("Connect", "Where should information go?", "CRM, email, booking, reporting, or follow-up workflow."),
            ("Refine", "What needs tuning?", "Better answers, prompts, scripts, and escalation rules."),
        ],
        widths=[1.2, 2.4, 2.9],
    )
    doc.add_page_break()

    doc.add_heading("7. Your AI Opportunity Map", level=1)
    add_para(
        doc,
        "Use this section as a working page. The goal is to identify where AI could create useful value without overcomplicating the business.",
    )
    add_table(
        doc,
        ["Area", "What happens now?", "Could AI help?", "Priority"],
        [
            ("Phone calls", "", "", ""),
            ("Website enquiries", "", "", ""),
            ("Quotes/bookings", "", "", ""),
            ("Customer questions", "", "", ""),
            ("Admin follow-up", "", "", ""),
            ("Staff knowledge", "", "", ""),
            ("Marketing/content", "", "", ""),
        ],
        widths=[1.35, 2.15, 2.15, 0.85],
    )
    add_callout(
        doc,
        "How to use this page",
        "Start with the area that is costing the most time or losing the most enquiries. A small project that fixes a real leak is better than a large project nobody maintains.",
    )
    doc.add_page_break()

    doc.add_heading("8. Self-Assessment Checklist", level=1)
    add_para(doc, "Tick the statements that are true for your business. The more boxes you tick, the more likely there is a practical AI opportunity worth exploring.")
    doc.add_heading("Customer response", level=2)
    add_checklist(
        doc,
        [
            "We miss calls when staff are busy, with customers, or out of hours.",
            "Customers ask the same questions before booking or buying.",
            "Our website does not answer enough practical questions.",
            "Our enquiry follow-up is slower than we would like.",
            "We rely heavily on customers calling because the website does not explain enough.",
        ],
    )
    doc.add_heading("Internal workload", level=2)
    add_checklist(
        doc,
        [
            "Staff repeat the same explanations in emails or calls.",
            "Quote or booking details are often incomplete.",
            "We manually copy information between systems.",
            "We do not have a clear customer knowledge base or FAQ.",
            "We spend too much time on admin that does not require judgement.",
        ],
    )
    doc.add_heading("Growth readiness", level=2)
    add_checklist(
        doc,
        [
            "We want more enquiries but our process is already stretched.",
            "We are paying for leads or ads but not capturing every opportunity.",
            "Competitors appear easier to understand or contact online.",
            "We want to use AI but need help choosing the first practical step.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("Your Score", level=1)
    add_para(doc, "Count how many boxes you ticked in the checklist.")
    add_table(
        doc,
        ["Score", "What it likely means", "Suggested next step"],
        [
            ("0-4", "AI may still help, but the first priority is probably clarity and process.", "Read the guide and identify one simple improvement."),
            ("5-9", "There are visible opportunities to save time or improve response.", "Consider an AI Website Review or Starter Pack."),
            ("10+", "AI could likely improve response, lead capture, and operations.", "Map voice, chat, website, and workflow opportunities."),
        ],
        widths=[1.1, 2.8, 2.6],
    )
    doc.add_page_break()

    doc.add_heading("Next Steps", level=1)
    add_para(
        doc,
        "You do not need to become an AI expert. You need to know where AI can help your business make the customer experience clearer, faster, or easier to manage.",
    )
    doc.add_heading("If you want to learn first", level=2)
    add_para(doc, "Use this guide to make a shortlist of the customer questions, missed moments, and admin tasks that matter most.")
    doc.add_heading("If your website may be leaking enquiries", level=2)
    add_para(doc, "Start with an AI Website Review. This checks your website clarity, lead capture, customer journey, AI search readiness, and practical automation opportunities.")
    doc.add_heading("If calls are the biggest issue", level=2)
    add_para(doc, "Start with an AI Voice Receptionist. This can answer calls, collect details, summarise conversations, and route enquiries to your team.")
    doc.add_heading("If customers need answers before calling", level=2)
    add_para(doc, "Start with a Website + AI Chatbot package. This creates a stronger digital front door and helps visitors understand what to do next.")
    add_callout(
        doc,
        "Final thought",
        "The best AI projects do not start with technology. They start with a clear customer problem, a practical first step, and a system that can be improved over time.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
